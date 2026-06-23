from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
OUT.mkdir(exist_ok=True)

NAVY = "17324D"
BLUE = "1D6FA5"
TEAL = "159A9C"
GOLD = "D7A441"
INK = "24303C"
MUTED = "64748B"
LIGHT = "EEF4F7"
PALE_BLUE = "E7F1F7"
PALE_GOLD = "FFF7E3"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"
GRID = "CAD6DF"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=None, color=INK, bold=None, italic=None, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color=GRID, size="12", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, 9, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def base_document(preset="standard"):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5 if preset == "compact" else 11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(5 if preset == "compact" else 6)
    normal.paragraph_format.line_spacing = 1.16 if preset == "standard" else 1.2

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in [
        ("Heading 1", 17, NAVY, 16, 8),
        ("Heading 2", 13.5, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name, left, first in [("List Bullet", 0.38, -0.18), ("List Number", 0.38, -0.18)]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(left)
        style.paragraph_format.first_line_indent = Inches(first)
        style.paragraph_format.space_after = Pt(3.5)
        style.paragraph_format.line_spacing = 1.2

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("WAKE BUSINESS SUITE  |  WAKE SERVICES")
    set_run(run, 8.5, MUTED, bold=True)
    paragraph_border_bottom(hp, GRID, "8", "4")

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)
    return doc


def add_cover(doc, kicker, title, subtitle, audience, version="Version 1.0"):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(kicker.upper())
    set_run(r, 10, TEAL, bold=True)
    p.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run(r, 14, BLUE)
    p.paragraph_format.space_after = Pt(36)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WAKE SERVICES")
    set_run(r, 13, NAVY, bold=True)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{audience}\n{version} · 22 juin 2026")
    set_run(r, 10.5, MUTED)
    p.paragraph_format.space_after = Pt(42)

    callout(doc, "Finalité", "Présenter une vision fiable du logiciel à partir de ses modules, parcours, rôles, règles de gestion et mécanismes de sécurité observés dans le code source.", PALE_BLUE)
    doc.add_page_break()


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_run(r, 8.5, TEAL, bold=True)
    return p


def callout(doc, label, text, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(label + "  ")
    set_run(r, 10.5, NAVY, bold=True)
    r = p.add_run(text)
    set_run(r, 10.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, header_fill=NAVY, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(header))
        set_run(r, font_size, WHITE, bold=True)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if row_index % 2 == 1:
                set_cell_shading(cells[idx], "F7FAFC")
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, font_size, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.38 + level * 0.24)
        p.add_run(item)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_toc_list(doc, items):
    table = doc.add_table(rows=0, cols=2)
    set_table_widths(table, [900, 8460])
    for number, label in items:
        cells = table.add_row().cells
        set_cell_shading(cells[0], PALE_BLUE)
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(number)
        set_run(r, 11, BLUE, bold=True)
        p = cells[1].paragraphs[0]
        r = p.add_run(label)
        set_run(r, 10.5, INK, bold=True)
    doc.add_paragraph()


def add_section_break(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_presentation():
    doc = base_document("standard")
    add_cover(
        doc,
        "Dossier de présentation",
        "WAKE Business Suite",
        "Plateforme intégrée de pilotage opérationnel et financier",
        "Document destiné à la Direction, aux responsables métiers et aux partenaires",
    )

    add_kicker(doc, "Vue d’ensemble")
    doc.add_heading("1. Synthèse exécutive", level=1)
    doc.add_paragraph(
        "WAKE Business Suite est une application web de gestion intégrée conçue pour centraliser les activités de WAKE SERVICES. "
        "Elle relie les processus commerciaux, financiers, de construction et de placement du personnel dans un environnement unique, "
        "avec gestion des rôles, notifications et traçabilité des actions."
    )
    callout(
        doc,
        "Proposition de valeur",
        "Un même référentiel pour décider, exécuter et contrôler : les opérations produisent des données immédiatement exploitables par la Direction.",
        PALE_GOLD,
    )
    add_table(
        doc,
        ["Dimension", "Apport du logiciel"],
        [
            ["Pilotage", "Tableau de bord Direction avec trésorerie, revenus, dépenses, marges, alertes et dossiers critiques."],
            ["Exécution", "Workflows guidés pour les demandes de fonds, ventes, livraisons, factures, chantiers et contrats de placement."],
            ["Contrôle", "Permissions par rôle, validation des opérations sensibles et journal d’audit."],
            ["Rentabilité", "Calculs de marges sur ventes, projets, contrats de placement et factures."],
            ["Traçabilité", "Références automatiques, statuts, historiques, pièces justificatives et notifications."],
        ],
        [1900, 7460],
    )

    doc.add_heading("2. Périmètre fonctionnel", level=1)
    add_table(
        doc,
        ["Domaine", "Fonctions principales", "Résultat attendu"],
        [
            ["Direction", "KPI, alertes, analyses détaillées, exports PDF/Excel", "Décisions plus rapides et suivi transversal"],
            ["Finance & trésorerie", "Demandes de fonds, comptes, mouvements, transferts, rapports", "Maîtrise des décaissements et soldes"],
            ["Commercial & ventes", "Clients, produits, devis, commandes, livraisons", "Cycle de vente structuré de bout en bout"],
            ["Facturation", "Factures multi-sources, paiements partiels, échéances, impression", "Vision consolidée des créances et encaissements"],
            ["Construction", "Projets, travaux, consommables, rapports journaliers, coûts, photos", "Pilotage physique et financier des chantiers"],
            ["Placement", "Agents, contrats, affectations, présences, facturation", "Suivi des effectifs placés et de la marge"],
            ["Administration", "Utilisateurs, rôles, permissions, audit, notifications", "Accès maîtrisés et responsabilité claire"],
        ],
        [1650, 4650, 3060],
        font_size=8.8,
    )

    doc.add_heading("3. Utilisateurs et responsabilités", level=1)
    add_table(
        doc,
        ["Rôle", "Responsabilités dominantes"],
        [
            ["Super Admin", "Administration complète, configuration des accès et supervision générale."],
            ["Direction", "Arbitrage, approbation des demandes, consultation des rapports et de l’audit."],
            ["Finance", "Facturation, encaissements, reporting financier et gestion de comptes autorisés."],
            ["Responsable Caisse/Banque", "Paiement des demandes approuvées et suivi des mouvements."],
            ["Chef de Projet", "Création et suivi des projets, rapports journaliers et demandes liées aux chantiers."],
            ["RH Placement", "Agents, contrats, présences et facturation du personnel placé."],
            ["Commercial", "Clients, produits, devis, commandes et suivi du portefeuille commercial."],
            ["Logistique", "Préparation des livraisons et suivi des quantités livrées."],
        ],
        [2600, 6760],
    )

    add_section_break(doc)
    add_kicker(doc, "Chaîne de valeur")
    doc.add_heading("4. Workflows métier structurants", level=1)
    doc.add_heading("Demande de fonds", level=2)
    add_steps(doc, [
        "Le demandeur saisit le besoin, le service, le montant, la devise, la date souhaitée et éventuellement un justificatif.",
        "La demande est enregistrée en brouillon ou soumise à la Direction.",
        "La Direction approuve en affectant un compte compatible avec la devise, ou rejette avec un motif obligatoire.",
        "Le responsable du compte confirme le paiement et peut joindre une preuve.",
        "Le système crée automatiquement le mouvement de trésorerie et conserve l’historique.",
    ])
    doc.add_heading("Cycle commercial", level=2)
    add_steps(doc, [
        "Création du client et du catalogue de produits ou prestations.",
        "Établissement du devis avec prix, coûts, taxes et marge estimée.",
        "Validation puis conversion du devis en commande.",
        "Livraison partielle ou totale selon les quantités réellement expédiées.",
        "Génération de facture, encaissements partiels ou complets et clôture.",
    ])
    doc.add_heading("Chantier", level=2)
    add_steps(doc, [
        "Création du projet avec contrat, budget prévisionnel, planning et chef de projet.",
        "Définition des travaux et consommables prévus.",
        "Saisie quotidienne des quantités exécutées, consommations, dépenses, blocages et photos.",
        "Calcul des avancements physique et financier, écarts de consommation, marge et retard.",
        "Consultation du cockpit et des rapports de construction.",
    ])
    doc.add_heading("Placement du personnel", level=2)
    add_steps(doc, [
        "Enregistrement des agents et de leurs informations professionnelles.",
        "Création du contrat client et affectation des agents avec coût et tarif client.",
        "Suivi mensuel des jours présents, absences et heures supplémentaires.",
        "Génération des factures de placement et contrôle de la marge.",
    ])

    doc.add_heading("5. Tableau de bord de Direction", level=1)
    doc.add_paragraph(
        "La page d’accueil est conçue comme un cockpit décisionnel. Les cartes sont interactives et ouvrent des analyses détaillées exportables."
    )
    add_bullets(doc, [
        "Soldes consolidés des caisses et banques.",
        "Demandes financières en attente et dépenses du mois.",
        "Revenus mensuels, factures impayées et marge globale estimée.",
        "Projets de construction actifs, avancement moyen et dossiers critiques.",
        "Personnel placé actif, commandes ouvertes et livraisons à traiter.",
        "Graphiques revenus/dépenses, structure des coûts et alertes prioritaires.",
    ])

    add_section_break(doc)
    add_kicker(doc, "Conception et contrôle")
    doc.add_heading("6. Architecture observée", level=1)
    doc.add_paragraph(
        "L’application suit une architecture MVC PHP légère : le routeur distribue les requêtes aux contrôleurs, les modèles accèdent à MySQL via PDO "
        "et les vues PHP rendent l’interface. Les permissions sont contrôlées au niveau des routes et de la navigation."
    )
    add_table(
        doc,
        ["Couche", "Composants observés"],
        [
            ["Interface", "Vues PHP, CSS responsive, JavaScript pour filtres, graphiques, modales et validations."],
            ["Application", "Contrôleurs métier, routeur, middlewares d’authentification et de permissions."],
            ["Données", "Modèles PHP, PDO, requêtes préparées et base MySQL structurée par migrations."],
            ["Sécurité", "Sessions, timeout, CSRF, échappement HTML, contrôle des uploads et journal d’audit."],
            ["Documents", "Exports PDF/Excel, impression de factures et gestion de pièces justificatives."],
        ],
        [1850, 7510],
    )

    doc.add_heading("7. Sécurité et gouvernance", level=1)
    add_bullets(doc, [
        "Authentification par email et mot de passe haché.",
        "Expiration automatique de session après 30 minutes d’inactivité.",
        "Contrôle d’accès fondé sur les permissions associées aux rôles.",
        "Jetons CSRF sur les formulaires sensibles.",
        "Requêtes PDO préparées et échappement des sorties HTML.",
        "Contrôle du type, de la taille et du format des fichiers téléversés.",
        "Journalisation des actions critiques et notifications internes.",
    ])
    callout(
        doc,
        "Principe de contrôle interne",
        "La séparation des rôles permet de distinguer la demande, l’approbation et le paiement, particulièrement dans le circuit des fonds.",
        PALE_BLUE,
    )

    doc.add_heading("8. Forces du logiciel", level=1)
    add_bullets(doc, [
        "Couverture fonctionnelle cohérente avec plusieurs métiers de WAKE SERVICES.",
        "Référentiel unifié des clients, opérations, factures et paiements.",
        "Workflows explicites, statuts lisibles et historique des décisions.",
        "Calcul de rentabilité présent dans les principaux modules.",
        "Interface orientée action avec filtres, alertes, tableaux et exports.",
        "Conception modulaire facilitant l’ajout de nouveaux domaines.",
    ])

    doc.add_heading("9. Points de vigilance avant production", level=1)
    add_table(
        doc,
        ["Priorité", "Constat", "Recommandation"],
        [
            ["Critique", "Le mode debug est activé dans la configuration.", "Désactiver APP_DEBUG et afficher des erreurs génériques en production."],
            ["Critique", "Les paramètres de base sont inscrits dans le fichier de configuration.", "Charger les secrets depuis des variables d’environnement ou un gestionnaire de secrets."],
            ["Élevée", "Un compte administrateur de démonstration est documenté.", "Changer immédiatement le mot de passe et supprimer les comptes de test."],
            ["Élevée", "La disponibilité dépend d’une base MySQL locale.", "Mettre en place sauvegardes automatiques, restauration testée et supervision."],
            ["Moyenne", "Les exports et uploads augmentent le volume de stockage.", "Définir rétention, quotas, antivirus et sauvegarde des pièces."],
            ["Moyenne", "La documentation technique et les tests automatisés sont limités.", "Ajouter tests de régression, procédure de déploiement et journal des versions."],
        ],
        [1200, 3550, 4610],
        header_fill="8B3A3A",
        font_size=8.6,
    )

    doc.add_heading("10. Feuille de route recommandée", level=1)
    add_table(
        doc,
        ["Horizon", "Actions"],
        [
            ["Avant mise en service", "Durcissement de la configuration, comptes nominatifs, sauvegardes, HTTPS, recette par rôle."],
            ["0–3 mois", "Formation des utilisateurs, référentiel de données, suivi des incidents et indicateurs d’adoption."],
            ["3–6 mois", "Tests automatisés, amélioration des exports, tableaux de bord par rôle et procédures d’archivage."],
            ["6–12 mois", "API/intégrations, notifications externes, analytique avancée et mobilité renforcée."],
        ],
        [1800, 7560],
    )
    callout(
        doc,
        "Conclusion",
        "WAKE Business Suite constitue une base solide d’ERP métier interne. Sa valeur dépendra surtout de la qualité des données, de la discipline des rôles et du durcissement de l’environnement de production.",
        PALE_GOLD,
    )

    path = OUT / "WAKE_Business_Suite_Dossier_de_Presentation.docx"
    doc.save(path)
    return path


def build_user_guide():
    doc = base_document("compact")
    add_cover(
        doc,
        "Guide utilisateur",
        "Utiliser WAKE Business Suite",
        "Parcours pratiques, procédures et règles de gestion",
        "Pour les utilisateurs opérationnels, responsables métiers et administrateurs",
    )

    add_kicker(doc, "Navigation")
    doc.add_heading("Sommaire", level=1)
    add_toc_list(doc, [
        ("01", "Premiers pas et navigation"),
        ("02", "Tableau de bord et outils transversaux"),
        ("03", "Commercial, ventes et livraisons"),
        ("04", "Facturation et encaissements"),
        ("05", "Finance, trésorerie et demandes de fonds"),
        ("06", "Construction et suivi de chantier"),
        ("07", "Placement du personnel"),
        ("08", "Administration, rôles et audit"),
        ("09", "Référentiel des statuts"),
        ("10", "Bonnes pratiques et dépannage"),
    ])

    doc.add_heading("1. Premiers pas", level=1)
    doc.add_heading("Se connecter", level=2)
    add_steps(doc, [
        "Ouvrez l’adresse communiquée par l’administrateur.",
        "Saisissez votre email professionnel et votre mot de passe.",
        "Cliquez sur « Se connecter ».",
        "Si un module n’apparaît pas, vérifiez avec l’administrateur que votre rôle possède la permission correspondante.",
    ])
    callout(doc, "Sécurité", "La session expire après 30 minutes d’inactivité. Enregistrez les formulaires longs avant de quitter votre poste.", PALE_GOLD)

    doc.add_heading("Comprendre l’interface", level=2)
    add_bullets(doc, [
        "Le menu gauche regroupe Pilotage, Commercial & Ventes, Finance & Trésorerie, Construction, Placement et Administration.",
        "Les éléments visibles dépendent de votre rôle et de vos permissions.",
        "Les boutons principaux créent ou confirment une opération ; les boutons secondaires permettent généralement de revenir, filtrer ou exporter.",
        "Les badges colorés indiquent le statut d’un dossier.",
        "Les tableaux proposent souvent une recherche instantanée et parfois un filtre par statut.",
        "Les notifications signalent les événements nécessitant une attention.",
    ])

    doc.add_heading("Règles générales de saisie", level=2)
    add_bullets(doc, [
        "Renseignez les champs marqués « Obligatoire » avant validation.",
        "Utilisez des libellés précis et évitez les abréviations ambiguës.",
        "Vérifiez la devise USD ou CDF avant de confirmer un montant.",
        "Joignez les justificatifs lisibles au format PDF, JPG ou PNG lorsque le formulaire le permet.",
        "Relisez le récapitulatif avant toute soumission, approbation, livraison ou paiement.",
    ])

    doc.add_heading("2. Tableau de bord", level=1)
    doc.add_paragraph(
        "Le tableau de bord présente la situation générale de l’entreprise. Cliquez sur une carte ou une zone d’analyse pour ouvrir le détail."
    )
    add_table(
        doc,
        ["Bloc", "Comment l’utiliser"],
        [
            ["KPI", "Contrôler les soldes, demandes en attente, dépenses, revenus, factures impayées, projets, placements et livraisons."],
            ["Graphiques", "Comparer revenus et dépenses et visualiser la structure des coûts."],
            ["Listes critiques", "Ouvrir les demandes récentes, projets sensibles et factures en retard."],
            ["Détail", "Cliquer sur une carte pour afficher les lignes sources."],
            ["Exports", "Dans la fenêtre de détail, utiliser Excel ou PDF selon le besoin."],
        ],
        [1900, 7460],
    )

    doc.add_heading("Notifications et rapports", level=2)
    add_bullets(doc, [
        "Ouvrez Notifications pour consulter les événements générés par les workflows.",
        "Marquez une notification comme lue après traitement.",
        "Ouvrez Rapports de gestion pour consolider les informations transversales.",
        "Utilisez les exports uniquement sur des données vérifiées.",
    ])

    add_section_break(doc)
    add_kicker(doc, "Commercial")
    doc.add_heading("3. Clients, produits, devis et commandes", level=1)
    doc.add_heading("Créer un client", level=2)
    add_steps(doc, [
        "Ouvrez Commercial & Ventes > Clients.",
        "Cliquez sur le bouton de création.",
        "Renseignez la raison sociale, le contact, le téléphone, l’email, l’adresse et les informations fiscales disponibles.",
        "Choisissez le statut Actif et enregistrez.",
    ])
    doc.add_heading("Créer un produit ou une prestation", level=2)
    add_steps(doc, [
        "Ouvrez Produits & stock puis créez une fiche.",
        "Saisissez le SKU, le nom, la catégorie, l’unité, le coût d’achat, le prix de vente et le stock.",
        "Vérifiez que le coût et le prix sont cohérents afin que la marge calculée soit fiable.",
        "Activez la fiche puis enregistrez.",
    ])
    doc.add_heading("Créer et convertir un devis", level=2)
    add_steps(doc, [
        "Ouvrez Devis puis « Créer un devis ».",
        "Sélectionnez le client, la date, la validité et ajoutez les notes commerciales.",
        "Pour chaque ligne, choisissez un produit ou une prestation libre, puis renseignez description, quantité, prix, coût et taxe.",
        "Créez le devis. Il est initialement en brouillon.",
        "Cliquez sur « Valider » après contrôle.",
        "Cliquez sur « Créer la commande » pour convertir le devis.",
    ])
    callout(doc, "Contrôle marge", "Le coût estimé n’est pas destiné au client ; il permet de calculer la marge commerciale. Ne le laissez pas à zéro sans justification.", PALE_BLUE)

    doc.add_heading("Traiter une commande et une livraison", level=2)
    add_steps(doc, [
        "Ouvrez Commandes et sélectionnez le dossier.",
        "Vérifiez les lignes, les quantités, le total et la marge.",
        "Cliquez sur « Préparer livraison ».",
        "Indiquez pour chaque article la quantité réellement livrée ; la livraison peut être partielle.",
        "Ajoutez les réserves ou instructions dans les notes et enregistrez.",
        "Depuis la commande, générez la facture lorsque les conditions de facturation sont remplies.",
    ])

    doc.add_heading("4. Facturation et paiements clients", level=1)
    doc.add_heading("Créer une facture manuelle", level=2)
    add_steps(doc, [
        "Ouvrez Finance & Trésorerie > Facturation puis « Créer une facture ».",
        "Sélectionnez le client et la source : autre service, projet de construction ou prestation diverse.",
        "Renseignez les dates, le statut Brouillon ou Envoyée, les conditions et les notes.",
        "Ajoutez les lignes avec description, quantité, prix unitaire, coût et taxe.",
        "Créez la facture et ouvrez-la pour vérifier le total, la marge et l’échéance.",
    ])
    doc.add_heading("Générer une facture métier", level=2)
    add_bullets(doc, [
        "Commande : utilisez « Générer facture » dans le dossier de commande.",
        "Placement : choisissez un contrat actif et le mois dans l’écran de création de facture.",
        "Construction : une facture peut être créée avec la source Projet construction et la référence utile en notes.",
    ])
    doc.add_heading("Enregistrer un encaissement", level=2)
    add_steps(doc, [
        "Ouvrez Paiements clients ou cliquez sur « Encaisser » depuis une facture.",
        "Sélectionnez la facture ouverte.",
        "Saisissez la date, le montant et la méthode : Cash, Banque, Mobile Money, Chèque ou Autre.",
        "Ajoutez la référence de transaction dans les notes.",
        "Enregistrez. Le reste à payer et le statut sont recalculés.",
    ])
    callout(doc, "Paiements partiels", "Une facture peut recevoir plusieurs encaissements. Ne saisissez jamais un montant supérieur au solde restant.", PALE_GOLD)

    add_section_break(doc)
    add_kicker(doc, "Finance")
    doc.add_heading("5. Demandes de fonds", level=1)
    doc.add_heading("Créer une demande", level=2)
    add_steps(doc, [
        "Ouvrez Demandes de fonds puis « Nouvelle demande ».",
        "Renseignez le titre, le service, la date souhaitée et une justification détaillée.",
        "Saisissez le montant et sélectionnez USD ou CDF.",
        "Ajoutez si possible un justificatif PDF, JPG ou PNG de 5 Mo maximum.",
        "Choisissez « Enregistrer comme brouillon » ou « Soumettre à la Direction ».",
    ])
    doc.add_heading("Suivre une demande", level=2)
    add_bullets(doc, [
        "Brouillon : le dossier attend sa soumission.",
        "En attente : la Direction doit statuer.",
        "Approuvée : un compte de trésorerie a été affecté ; le paiement reste à exécuter.",
        "Rejetée : le motif est visible dans le dossier.",
        "Payée : le mouvement de trésorerie et la preuve éventuelle sont conservés.",
    ])
    doc.add_heading("Approuver ou rejeter — Direction", level=2)
    add_steps(doc, [
        "Ouvrez une demande en attente et cliquez sur l’action d’approbation.",
        "Lisez la justification et ouvrez les pièces jointes.",
        "Pour approuver, sélectionnez un compte actif dans la même devise et contrôlez son solde.",
        "Ajoutez éventuellement une instruction puis confirmez.",
        "Pour rejeter, saisissez un motif explicite et confirmez.",
    ])
    doc.add_heading("Payer — Responsable du compte", level=2)
    add_steps(doc, [
        "Ouvrez une demande approuvée.",
        "Vérifiez le montant, le compte affecté, le solde et le responsable.",
        "Saisissez la description du paiement et joignez éventuellement une preuve.",
        "Confirmez le paiement. Le système crée le mouvement et met à jour le solde.",
    ])

    doc.add_heading("6. Caisses, banques, mouvements et transferts", level=1)
    doc.add_heading("Comptes de trésorerie", level=2)
    add_bullets(doc, [
        "Créez un compte en précisant le type, la devise, le solde initial, le responsable et le statut.",
        "Consultez le détail d’un compte pour voir le solde et l’historique des mouvements.",
        "N’inactivez pas un compte impliqué dans des opérations non terminées sans validation Finance.",
    ])
    doc.add_heading("Transfert entre comptes", level=2)
    add_steps(doc, [
        "Créez le transfert avec compte source, compte destination, montants et motif.",
        "Enregistrez en brouillon ou soumettez.",
        "La personne autorisée approuve ou rejette.",
        "L’exécutant confirme le transfert ; les mouvements correspondants sont alors enregistrés.",
    ])
    callout(doc, "Devises", "Un transfert peut distinguer montant source et montant destination. Vérifiez le taux ou la justification lorsqu’ils diffèrent.", PALE_BLUE)

    add_section_break(doc)
    add_kicker(doc, "Construction")
    doc.add_heading("7. Projets et rapports journaliers", level=1)
    doc.add_heading("Créer un projet", level=2)
    add_steps(doc, [
        "Ouvrez Construction > Projets puis créez un projet.",
        "Renseignez le nom, le client, le montant du contrat, le coût prévisionnel, les dates et la localisation.",
        "Affectez un chef de projet et choisissez le statut.",
        "Ajoutez les travaux prévus avec unités, quantités, coûts et durées.",
        "Ajoutez les consommables prévus avec quantités et coûts.",
        "Enregistrez puis ouvrez le cockpit du projet.",
    ])
    doc.add_heading("Lire le cockpit chantier", level=2)
    add_table(
        doc,
        ["Indicateur", "Interprétation"],
        [
            ["Avancement global", "Progression physique pondérée des travaux."],
            ["Budget consommé", "Part du coût prévisionnel déjà engagée."],
            ["Écart consommation", "Différence entre consommations prévues et réelles."],
            ["Marge réelle estimée", "Montant du contrat diminué des coûts enregistrés."],
            ["Retard", "Écart en jours par rapport à la date de fin prévue."],
        ],
        [2450, 6910],
    )
    doc.add_heading("Saisir un rapport journalier", level=2)
    add_steps(doc, [
        "Depuis le projet, cliquez sur « Nouveau rapport journalier ».",
        "Renseignez la date, la météo, les observations et les blocages.",
        "Pour chaque travail, décrivez l’exécution, la quantité réalisée et le pourcentage d’avancement.",
        "Ajoutez les consommables utilisés et leurs coûts.",
        "Ajoutez les dépenses du jour.",
        "Joignez les photos chantier et une légende, puis enregistrez.",
    ])
    callout(doc, "Qualité des données", "Le cockpit dépend directement des quantités, coûts et pourcentages saisis. Corrigez les incohérences le jour même.", PALE_GOLD)

    add_section_break(doc)
    add_kicker(doc, "Placement")
    doc.add_heading("8. Agents, contrats et présences", level=1)
    doc.add_heading("Créer un agent", level=2)
    add_steps(doc, [
        "Ouvrez Placement > Agents puis créez une fiche.",
        "Renseignez l’identité, le poste, les contacts, le salaire de base et le statut.",
        "Enregistrez puis utilisez l’agent lors de la création d’un contrat.",
    ])
    doc.add_heading("Créer un contrat de placement", level=2)
    add_steps(doc, [
        "Ouvrez Contrats clients puis « Créer un contrat ».",
        "Saisissez le client, le contact, la période, le jour de facturation et le statut.",
        "Affectez les agents un par un.",
        "Pour chaque affectation, saisissez le poste, le coût agent, le tarif client et la date de début.",
        "Vérifiez que le tarif client est supérieur au coût lorsque le contrat doit générer une marge positive.",
        "Enregistrez et ouvrez le contrat pour contrôler la marge mensuelle.",
    ])
    doc.add_heading("Saisir les présences", level=2)
    add_steps(doc, [
        "Ouvrez Présences et sélectionnez le mois.",
        "Choisissez l’agent placé.",
        "Saisissez les jours présents, les absences et les heures supplémentaires.",
        "Ajoutez une observation en cas d’écart, puis enregistrez.",
    ])
    doc.add_heading("Générer la facture de placement", level=2)
    add_steps(doc, [
        "Ouvrez Factures placement ou la facturation centralisée.",
        "Sélectionnez le contrat et la période.",
        "Lancez la génération.",
        "Contrôlez le chiffre d’affaires, le coût et la marge avant envoi au client.",
    ])

    add_section_break(doc)
    add_kicker(doc, "Administration")
    doc.add_heading("9. Utilisateurs, rôles et audit", level=1)
    doc.add_heading("Créer un utilisateur", level=2)
    add_steps(doc, [
        "Ouvrez Administration > Utilisateurs.",
        "Créez la fiche avec nom, email, rôle, statut et mot de passe d’au moins 8 caractères.",
        "Utilisez un compte nominatif par personne ; ne partagez pas les identifiants.",
        "Désactivez immédiatement les comptes des personnes qui quittent l’organisation.",
    ])
    doc.add_heading("Gérer les permissions", level=2)
    add_bullets(doc, [
        "Ouvrez Rôles & permissions.",
        "Accordez uniquement les permissions nécessaires à la fonction.",
        "Séparez autant que possible création, approbation et paiement.",
        "Testez le rôle avec un compte dédié avant mise à disposition.",
    ])
    doc.add_heading("Consulter le journal d’audit", level=2)
    add_bullets(doc, [
        "Filtrez par utilisateur, action ou adresse IP.",
        "Utilisez l’audit pour reconstituer une décision ou une opération sensible.",
        "Ne modifiez pas directement les données en base pour contourner un workflow.",
    ])

    doc.add_heading("10. Référentiel des statuts", level=1)
    add_table(
        doc,
        ["Objet", "Statuts principaux", "Sens"],
        [
            ["Demande de fonds", "Draft, Pending, Approved, Rejected, Paid, Cancelled", "Brouillon, en attente, approuvée, rejetée, payée ou annulée."],
            ["Transfert", "Draft, Pending, Approved, Executed, Rejected, Cancelled", "Préparation, validation puis exécution."],
            ["Facture", "Draft, Sent, Partially Paid, Paid, Overdue, Cancelled", "Cycle d’émission et de recouvrement."],
            ["Projet", "Planning, In Progress, On Hold, Completed, Cancelled", "Planifié, en cours, suspendu, terminé ou annulé."],
            ["Contrat placement", "Draft, Active, Suspended, Expired, Closed", "Préparation, activité, suspension ou clôture."],
            ["Fiches", "active, inactive", "Disponible ou désactivée."],
        ],
        [1900, 3300, 4160],
        font_size=8.8,
    )

    doc.add_heading("11. Bonnes pratiques", level=1)
    add_bullets(doc, [
        "Travaillez avec un compte personnel et verrouillez votre poste en cas d’absence.",
        "Vérifiez toujours le client, la devise, le montant, la période et le compte avant confirmation.",
        "Conservez des justificatifs lisibles et nommés clairement.",
        "Ne dupliquez pas un client, un produit ou une facture sans recherche préalable.",
        "Utilisez les notes pour expliquer une exception ou une décision.",
        "Traitez régulièrement les alertes, notifications, factures en retard et demandes en attente.",
        "Signalez immédiatement toute erreur de solde ou opération non reconnue.",
    ])

    doc.add_heading("12. Dépannage rapide", level=1)
    add_table(
        doc,
        ["Situation", "Action recommandée"],
        [
            ["Impossible de se connecter", "Vérifier l’email, le clavier, le mot de passe et le statut du compte ; contacter l’administrateur."],
            ["Module absent du menu", "Le rôle ne possède probablement pas la permission nécessaire."],
            ["Formulaire refusé", "Contrôler les champs obligatoires, formats, montants positifs et dates."],
            ["Fichier refusé", "Utiliser PDF/JPG/PNG, respecter la taille autorisée et éviter les noms inhabituels."],
            ["Session expirée", "Reconnectez-vous puis reprenez l’opération ; évitez les longues périodes sans enregistrement."],
            ["Solde incohérent", "Ne pas corriger directement en base ; vérifier mouvements, transferts et paiements puis alerter Finance."],
            ["Facture non soldée", "Comparer le total des paiements au reste à payer et rechercher un encaissement manquant."],
        ],
        [2650, 6710],
        font_size=9.0,
    )
    callout(
        doc,
        "Assistance",
        "Lors d’un signalement, communiquez votre nom, le module, la référence du dossier, l’heure approximative et le message affiché. Ne transmettez jamais votre mot de passe.",
        PALE_RED,
    )

    path = OUT / "WAKE_Business_Suite_Guide_Utilisateur.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    presentation = build_presentation()
    guide = build_user_guide()
    print(presentation)
    print(guide)
